import streamlit as st
from dotenv import load_dotenv
load_dotenv()

from parser import extract_text
from analyzer import analyze_contract, compare_contracts, rewrite_contract
from formatter import format_results, get_risk_label

import hashlib
import base64
from io import BytesIO
from docx import Document

st.set_page_config(page_title="Dracu-Law", page_icon="⚖️", layout="wide")

# ---------- AUTH GUARD ----------
if not st.session_state.get("logged_in", False):
    st.switch_page("landing.py")

# ---------- HANDLE QUERY PARAM NAV ----------
if st.query_params.get("goto") == "home":
    st.query_params.clear()
    st.switch_page("landing.py")

if st.query_params.get("logout") == "true":
    st.session_state.logged_in = False
    st.session_state.username = ""
    st.query_params.clear()
    st.switch_page("landing.py")

# ---------- LOGO ----------
def get_base64_image(path):
    try:
        with open(path, "rb") as f:
            return base64.b64encode(f.read()).decode()
    except:
        return None

logo = get_base64_image("logo.png")

# ---------- GLOBAL CSS ----------
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:ital,wght@0,700;0,900;1,700&family=DM+Sans:wght@300;400;500;600&family=Courier+Prime&display=swap');

#MainMenu, footer, header, [data-testid="stToolbar"], .stDeployButton {
    visibility: hidden !important; display: none !important;
}

/* ── BASE ── */
body, .stApp {
    background: #0a0a0b !important;
    color: #e8e4dc !important;
    font-family: 'DM Sans', sans-serif !important;
}
.stApp::before {
    content: '';
    position: fixed; top: 0; left: 0; right: 0; height: 340px;
    background: radial-gradient(ellipse 80% 55% at 50% 0%, rgba(192,57,43,0.14) 0%, transparent 70%);
    pointer-events: none; z-index: 0;
}

/* ── PAGE TITLE ── */
.page-title { text-align: center; padding: 30px 24px 6px; }
.page-title h1 {
    font-family: 'Playfair Display', serif;
    font-size: clamp(26px, 4vw, 42px);
    font-weight: 900; color: #fff;
    letter-spacing: -1.5px; margin: 0;
}
.page-title p { font-size: 15px; color: #888; margin-top: 8px; }

/* ── SIDEBAR ── */
[data-testid="stSidebar"] { background: #0f0f11 !important; border-right: 1px solid #222 !important; }
[data-testid="stSidebar"] label, [data-testid="stSidebar"] p,
[data-testid="stSidebar"] span, [data-testid="stSidebar"] .stMarkdown { color: #bbb !important; }
[data-testid="stSidebar"] h1, [data-testid="stSidebar"] h2, [data-testid="stSidebar"] h3 {
    color: #eee !important; font-family: 'Playfair Display', serif !important;
}
[data-testid="stSidebar"] .stButton > button {
    background: #1a1a1e !important; border: 1px solid #2a2a2e !important;
    color: #bbb !important; width: 100% !important; padding: 10px !important;
    border-radius: 8px !important; font-size: 14px !important; font-weight: 500 !important;
}
[data-testid="stSidebar"] .stButton > button:hover {
    background: rgba(192,57,43,0.14) !important; border-color: rgba(192,57,43,0.4) !important;
    color: #fff !important; box-shadow: none !important;
}
[data-testid="stSidebar"] .stSelectbox > div > div {
    background: #1a1a1e !important; border-color: #2a2a2e !important; color: #eee !important;
}

/* ── TABS ── */
div[data-testid="stTabs"] { margin-top: 12px; }
div[data-testid="stTabs"] button {
    font-family: 'DM Sans', sans-serif !important;
    font-size: 14px !important; font-weight: 600 !important; color: #888 !important;
    padding: 10px 20px !important; background: transparent !important; border: none !important;
}
div[data-testid="stTabs"] button[aria-selected="true"] {
    color: #e74c3c !important; border-bottom: 2px solid #e74c3c !important;
}
div[data-testid="stTabs"] [data-testid="stTabsContent"] {
    background: #0d0d0f; border: 1px solid #1e1e22;
    border-radius: 0 0 12px 12px; padding: 28px 24px;
}

/* ── FILE UPLOADER ── */
[data-testid="stFileUploader"] {
    background: transparent !important;
    border-radius: 12px !important;
}
[data-testid="stFileUploader"] > section {
    background: #131316 !important;
    border: 2px solid #e07a3a !important;
    border-radius: 12px !important;
}
[data-testid="stFileUploaderDropzone"] {
    background: #131316 !important;
    border: 2px solid #e07a3a !important;
    border-radius: 12px !important;
}
[data-testid="stFileUploaderDropzone"]:hover,
[data-testid="stFileUploader"] > section:hover {
    border-color: #ff8c45 !important;
    background: #17130f !important;
    box-shadow: 0 0 16px rgba(224,122,58,0.2) !important;
}
[data-testid="stFileUploader"] label,
[data-testid="stFileUploader"] span,
[data-testid="stFileUploader"] p,
[data-testid="stFileUploader"] small,
[data-testid="stFileUploaderDropzone"] span,
[data-testid="stFileUploaderDropzone"] p,
[data-testid="stFileUploaderDropzone"] small {
    color: #ffffff !important;
    background: transparent !important;
}
[data-testid="stFileUploader"] > label {
    color: #e8e4dc !important;
    font-family: 'DM Sans', sans-serif !important;
    font-size: 14px !important;
    font-weight: 500 !important;
}
[data-testid="stFileUploaderDropzone"] button,
[data-testid="stFileUploader"] section button {
    background: rgba(224,122,58,0.12) !important;
    border: 1px solid rgba(224,122,58,0.5) !important;
    color: #f0944a !important;
    border-radius: 7px !important;
    font-size: 13px !important;
    font-weight: 600 !important;
    padding: 7px 18px !important;
    font-family: 'DM Sans', sans-serif !important;
    width: auto !important;
    min-width: unset !important;
}
[data-testid="stFileUploaderDropzone"] button:hover,
[data-testid="stFileUploader"] section button:hover {
    background: rgba(224,122,58,0.25) !important;
    border-color: #e07a3a !important;
    color: #fff !important;
    box-shadow: 0 0 12px rgba(224,122,58,0.3) !important;
    transform: none !important;
}
[data-testid="stFileUploaderFile"] {
    background: #1c1c20 !important;
    border: 1px solid #333 !important;
    border-radius: 8px !important;
}
[data-testid="stFileUploaderFile"] span,
[data-testid="stFileUploaderFile"] p { color: #e0dbd2 !important; }
[data-testid="stFileUploaderDeleteBtn"] button {
    background: transparent !important; border: none !important; color: #666 !important;
    width: auto !important; min-width: unset !important;
}
[data-testid="stFileUploaderDeleteBtn"] button:hover {
    color: #e74c3c !important; box-shadow: none !important;
    transform: none !important; background: transparent !important;
}

/* ── ALL MAIN BUTTONS ── */
.stButton > button {
    background: #d4601e !important;
    background-image: linear-gradient(180deg, #e8722a 0%, #c4581a 100%) !important;
    border-top: 1px solid #f0844a !important;
    border-left: 1px solid #c85010 !important;
    border-right: 1px solid #c85010 !important;
    border-bottom: 4px solid #7a3008 !important;
    border-radius: 8px !important;
    padding: 11px 32px 12px !important;
    min-width: 160px !important;
    width: auto !important;
    color: #ffffff !important;
    font-family: 'DM Sans', sans-serif !important;
    font-weight: 700 !important;
    font-size: 15px !important;
    letter-spacing: 0.3px !important;
    text-shadow: 0 1px 2px rgba(0,0,0,0.5) !important;
    box-shadow: 0 2px 0 rgba(0,0,0,0.4), 0 6px 24px rgba(200,80,20,0.45) !important;
    cursor: pointer !important;
    transition: all 0.1s ease !important;
    position: relative !important;
    top: 0 !important;
}
.stButton > button p,
.stButton > button span {
    color: #ffffff !important;
    font-weight: 700 !important;
    font-size: 15px !important;
    text-shadow: 0 1px 2px rgba(0,0,0,0.5) !important;
}
.stButton > button:hover {
    background: #e06e28 !important;
    background-image: linear-gradient(180deg, #f07e38 0%, #d06820 100%) !important;
    border-top-color: #ffaa70 !important;
    border-bottom-color: #6a2800 !important;
    top: -1px !important;
    box-shadow: 0 3px 0 rgba(0,0,0,0.4), 0 10px 30px rgba(200,80,20,0.6) !important;
}
.stButton > button:active {
    background-image: linear-gradient(180deg, #c05010 0%, #d46020 100%) !important;
    border-top-color: #a04010 !important;
    border-bottom-width: 1px !important;
    top: 3px !important;
    box-shadow: 0 1px 6px rgba(200,80,20,0.3) !important;
}

/* ── SIDEBAR BUTTON OVERRIDE ── */
[data-testid="stSidebar"] .stButton > button {
    background: #1a1a1e !important;
    background-image: none !important;
    border: 1px solid #2a2a2e !important;
    border-bottom: 1px solid #2a2a2e !important;
    color: #bbb !important;
    font-weight: 500 !important;
    font-size: 14px !important;
    padding: 10px !important;
    min-width: unset !important;
    box-shadow: none !important;
    text-shadow: none !important;
    top: 0 !important;
}
[data-testid="stSidebar"] .stButton > button p,
[data-testid="stSidebar"] .stButton > button span {
    color: #bbb !important;
    font-weight: 500 !important;
    text-shadow: none !important;
}
[data-testid="stSidebar"] .stButton > button:hover {
    background: rgba(192,57,43,0.14) !important;
    background-image: none !important;
    border-color: rgba(192,57,43,0.4) !important;
    border-bottom-color: rgba(192,57,43,0.4) !important;
    color: #fff !important;
    box-shadow: none !important;
    top: 0 !important;
}

/* ── DOWNLOAD BUTTON ── */
[data-testid="stDownloadButton"] > button {
    background: #1a4a7a !important;
    background-image: linear-gradient(180deg, #2060a0 0%, #163d6a 100%) !important;
    border-top: 1px solid #3a80cc !important;
    border-left: 1px solid #1a4a7a !important;
    border-right: 1px solid #1a4a7a !important;
    border-bottom: 4px solid #0a2040 !important;
    color: #a8d8ff !important;
    border-radius: 8px !important;
    font-size: 14px !important;
    font-weight: 600 !important;
    padding: 10px 24px 11px !important;
    width: auto !important;
    min-width: unset !important;
    text-shadow: 0 1px 2px rgba(0,0,0,0.5) !important;
    box-shadow: 0 4px 18px rgba(20,80,160,0.45) !important;
    cursor: pointer !important;
    transition: all 0.1s ease !important;
    position: relative !important;
    top: 0 !important;
}
[data-testid="stDownloadButton"] > button:hover {
    background-image: linear-gradient(180deg, #2a70ba 0%, #1c4e8a 100%) !important;
    border-top-color: #60a8ff !important;
    top: -1px !important;
    box-shadow: 0 6px 24px rgba(20,80,160,0.6) !important;
    color: #cceeff !important;
}
[data-testid="stDownloadButton"] > button:active {
    border-bottom-width: 1px !important;
    top: 3px !important;
    box-shadow: 0 1px 6px rgba(20,80,160,0.3) !important;
}
[data-testid="stDownloadButton"] > button p {
    color: #a8d8ff !important;
    font-weight: 600 !important;
}

/* ── ALERTS ── */
.stAlert { border-radius: 9px !important; }
div[data-baseweb="notification"] { border-radius: 9px !important; }

/* ── HEADINGS ── */
h2, h3, .stSubheader {
    font-family: 'Playfair Display', serif !important;
    color: #f0ece4 !important; letter-spacing: -0.5px !important;
}

/* ── PROGRESS BAR ── */
[data-testid="stProgressBar"] > div > div {
    background: linear-gradient(90deg, #c0392b, #e74c3c) !important;
    border-radius: 4px !important;
}
[data-testid="stProgressBar"] > div {
    background: #1e1e22 !important; border-radius: 4px !important;
}

/* ── TEXT AREA ── */
textarea {
    background: #16161a !important; border: 1px solid #2a2a2e !important;
    border-radius: 9px !important; color: #e0dbd2 !important;
    font-family: 'DM Sans', sans-serif !important;
}

/* ── CHECKBOXES ── */
.stCheckbox label,
.stCheckbox label p,
.stCheckbox span,
[data-testid="stCheckbox"] label,
[data-testid="stCheckbox"] span {
    color: #e8e4dc !important;
    font-size: 14px !important;
}

/* ── DIVIDER ── */
hr { border-color: #2a2a38 !important; }

/* ── FOOTER ── */
.dracu-footer {
    text-align: center; padding: 30px;
    border-top: 1px solid #1e1e22;
    color: #666; font-size: 13px; margin-top: 20px;
}
.dracu-footer strong { color: #888; }
</style>
""", unsafe_allow_html=True)

# ---------- NAVBAR ----------
username = st.session_state.get("username", "User")
logo_html = f'<img src="data:image/png;base64,{logo}" style="height:46px;margin-right:10px;vertical-align:middle;">' if logo else "⚖️"

nav_left, nav_right = st.columns([3, 1])

with nav_left:
    st.markdown(f"""
    <div style="display:flex;align-items:center;padding:10px 0px;height:68px;">
        {logo_html}
        <div>
            <div style="font-family:'Playfair Display',serif;font-size:24px;font-weight:900;color:#fff;letter-spacing:-0.5px;">
                Dracu<span style="color:#e74c3c;">-Law</span>
            </div>
            <div style="font-size:10px;color:#7a7a8a;letter-spacing:1.4px;text-transform:uppercase;margin-top:1px;">
                AI Legal Intelligence Engine
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

with nav_right:
    # Pure HTML links — completely bypasses Streamlit button styling
    st.markdown(f"""
    <div style="display:flex;align-items:center;justify-content:flex-end;gap:8px;height:68px;">
        <span style="background:rgba(192,57,43,0.12);border:1px solid rgba(192,57,43,0.3);
            border-radius:100px;padding:5px 12px;font-size:12px;color:#f07a6a;
            font-weight:600;white-space:nowrap;">
            👤 {username}
        </span>
        <a href="/?goto=home" target="_self" style="
            display:inline-block;
            background:transparent;
            border:1px solid #3a3a48;
            border-radius:5px;
            color:#b0a898;
            font-size:11px;
            font-weight:500;
            padding:5px 10px;
            text-decoration:none;
            white-space:nowrap;
            font-family:'DM Sans',sans-serif;
            line-height:1.4;
            cursor:pointer;
        " onmouseover="this.style.background='rgba(255,255,255,0.06)';this.style.color='#e0dbd2';this.style.borderColor='#555560';"
           onmouseout="this.style.background='transparent';this.style.color='#b0a898';this.style.borderColor='#3a3a48';">
            🏠 Home
        </a>
        <a href="/?logout=true" target="_self" style="
            display:inline-block;
            background:transparent;
            border:1px solid #3a3a48;
            border-radius:5px;
            color:#b0a898;
            font-size:11px;
            font-weight:500;
            padding:5px 10px;
            text-decoration:none;
            white-space:nowrap;
            font-family:'DM Sans',sans-serif;
            line-height:1.4;
            cursor:pointer;
        " onmouseover="this.style.background='rgba(255,255,255,0.06)';this.style.color='#e0dbd2';this.style.borderColor='#555560';"
           onmouseout="this.style.background='transparent';this.style.color='#b0a898';this.style.borderColor='#3a3a48';">
            Log Out
        </a>
    </div>
    """, unsafe_allow_html=True)

st.markdown("<hr style='border-color:#2a2a2e;margin:0 0 8px 0;'>", unsafe_allow_html=True)

# ---------- PAGE TITLE ----------
st.markdown("""
<div class="page-title">
    <h1>AI Contract Intelligence Platform</h1>
    <p>Upload, analyze, compare, and rewrite contracts with surgical precision.</p>
</div>
""", unsafe_allow_html=True)

st.divider()

# ---------- SIDEBAR ----------
ALLOWED_MODELS = [
    "llama-3.3-70b-versatile",
    "openai/gpt-oss-120b",
    "openai/gpt-oss-20b",
    "llama-3.1-8b-instant"
]

with st.sidebar:
    st.markdown("### ⚙️ Settings")
    model_choice = st.selectbox("Model", ALLOWED_MODELS)
    if model_choice not in ALLOWED_MODELS:
        st.error("Invalid model selected.")
        st.stop()
    st.caption("Powered by Groq LLMs")
    st.divider()
    st.markdown("**🔄 Session**")
    if st.button("Clear & Start Over", key="btn_clear"):
        keys_to_keep = {"logged_in", "username", "user_email"}
        for k in [k for k in list(st.session_state.keys()) if k not in keys_to_keep]:
            del st.session_state[k]
        st.success("Session cleared!")
        st.rerun()

# ---------- HELPERS ----------
def get_file_hash(uploaded_file):
    uploaded_file.seek(0)
    content = uploaded_file.read()
    uploaded_file.seek(0)
    return hashlib.md5(content).hexdigest()

def create_doc(text):
    doc = Document()
    doc.add_heading("Improved Contract", 0)
    p = doc.add_paragraph()
    run = p.add_run("Generated by Dracu-Law AI\n")
    run.italic = True
    doc.add_paragraph("")
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
        elif stripped.isupper() and len(stripped) > 3:
            doc.add_heading(stripped, level=2)
        elif stripped.endswith(":") and len(stripped) < 60:
            doc.add_heading(stripped.rstrip(":"), level=3)
        else:
            doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def render_risk_display(score):
    if score is None:
        st.warning("Risk score unavailable.")
        return
    if score <= 3:
        color   = "#27ae60"
        bg      = "#0b1a10"
        border  = "#1a4a28"
        icon    = "✅"
        verdict = "Low Risk — generally safe to sign"
    elif score <= 6:
        color   = "#e6a817"
        bg      = "#1a1608"
        border  = "#4a3c10"
        icon    = "⚠️"
        verdict = "Moderate Risk — review flagged clauses carefully"
    else:
        color   = "#e74c3c"
        bg      = "#1a0b0b"
        border  = "#4a1818"
        icon    = "🚨"
        verdict = "High Risk — seek legal advice before signing"

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown(
            f"<div style='background:{bg};border:2px solid {border};border-radius:14px;"
            f"padding:28px 28px;text-align:center;margin-bottom:12px;'>"
            f"<div style='font-size:40px;margin-bottom:8px;'>{icon}</div>"
            f"<div style='font-family:\"Playfair Display\",serif;font-size:58px;font-weight:900;"
            f"color:{color};line-height:1;'>{score}"
            f"<span style='font-size:22px;color:#444;'>/10</span></div>"
            f"<div style='font-size:15px;color:{color};font-weight:600;margin-top:10px;"
            f"letter-spacing:0.2px;'>{verdict}</div>"
            f"</div>",
            unsafe_allow_html=True
        )
        st.progress(min(max(score / 10, 0.0), 1.0))

# ---------- TABS ----------
tab1, tab2, tab3 = st.tabs([
    "🔬  Analyze Contract",
    "⚖️  Compare Contracts",
    "✍️  Improve Contract"
])

# =========================================================
# TAB 1 — ANALYZE
# =========================================================
with tab1:
    uploaded_file = st.file_uploader("Upload contract PDF", type=["pdf"], key="single")

    if uploaded_file:
        st.success("✅ File uploaded successfully.")
        file_hash = get_file_hash(uploaded_file)

        if st.button("🔬 Analyze Contract", key="btn_analyze"):
            for key in ["last_contract", "last_flags", "last_hash", "last_result"]:
                st.session_state.pop(key, None)
            try:
                with st.spinner("Extracting text from PDF..."):
                    text = extract_text(uploaded_file)
                with st.spinner("Analyzing contract with AI (this may take 10–30 seconds)..."):
                    raw = analyze_contract(text, model_choice)
                result = format_results(raw)
                st.session_state.last_contract = text
                st.session_state.last_flags    = result["red_flags"]
                st.session_state.last_hash     = file_hash
                st.session_state.last_result   = result
            except Exception as e:
                st.error(f"Analysis failed: {e}")
                st.stop()

        if st.session_state.get("last_hash") == file_hash and "last_result" in st.session_state:
            result = st.session_state.last_result
            st.subheader("Risk Level")
            render_risk_display(result.get("risk_score"))
            st.divider()

            c1, c2 = st.columns(2)
            with c1:
                st.subheader("🚩 Red Flags")
                for flag in result["red_flags"]:
                    st.warning(flag)
            with c2:
                st.subheader("💡 Suggestions")
                for suggestion in result["suggestions"]:
                    st.info(suggestion)

            st.subheader("📋 Summary")
            st.write(result["explanation"])

# =========================================================
# TAB 2 — COMPARE
# =========================================================
with tab2:
    f1 = st.file_uploader("Upload Contract A", type=["pdf"], key="a")
    f2 = st.file_uploader("Upload Contract B", type=["pdf"], key="b")

    if f1 and not f2:
        st.warning("📂 Please upload **Contract B** to continue.")
    elif f2 and not f1:
        st.warning("📂 Please upload **Contract A** to continue.")

    if f1 and f2:
        if st.button("⚖️ Compare Contracts", key="btn_compare"):
            try:
                with st.spinner("Extracting text from both contracts..."):
                    t1 = extract_text(f1)
                    t2 = extract_text(f2)
                with st.spinner("Comparing contracts with AI..."):
                    verdict = compare_contracts(t1, t2, model_choice)
                st.session_state.compare_verdict = verdict
            except Exception as e:
                st.error(f"Comparison failed: {e}")

        if "compare_verdict" in st.session_state:
            verdict = st.session_state.compare_verdict
            lines = verdict.strip().split("\n")
            winner_line     = ""
            reason_line     = ""
            differences     = []
            current_section = None

            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if line.upper().startswith("WINNER"):
                    winner_line = line.split(":", 1)[-1].strip()
                    current_section = "winner"
                elif line.upper().startswith("REASON"):
                    reason_line = line.split(":", 1)[-1].strip()
                    current_section = "reason"
                elif "KEY DIFF" in line.upper():
                    current_section = "diff"
                elif current_section == "reason" and not reason_line:
                    reason_line = line
                elif current_section == "diff" and line.startswith(("-", "\u2022", "*")):
                    differences.append(line.lstrip("-\u2022* ").strip())
                elif current_section == "diff" and line:
                    differences.append(line.strip())

            st.divider()
            st.markdown("### ⚖️ Contract Comparison Verdict")
            st.caption("AI ANALYSIS")
            st.divider()

            winner_text = winner_line or "See full analysis below"
            st.markdown("**🏆 WINNER**")
            st.success(f"### {winner_text}")

            if reason_line:
                st.markdown("")
                st.markdown("**📋 REASON**")
                st.info(reason_line)

            if differences:
                st.markdown("")
                st.markdown("**🔍 KEY DIFFERENCES**")
                st.markdown("")
                for i, item in enumerate(differences, 1):
                    col_num, col_text = st.columns([0.06, 0.94])
                    with col_num:
                        st.markdown(
                            f"<div style='background:#c0392b;color:white;border-radius:50%;"
                            f"width:28px;height:28px;display:flex;align-items:center;"
                            f"justify-content:center;font-size:12px;font-weight:700;"
                            f"margin-top:4px;'>{i}</div>",
                            unsafe_allow_html=True
                        )
                    with col_text:
                        st.markdown(item)
                    st.markdown("---")

# =========================================================
# TAB 3 — IMPROVE CONTRACT
# =========================================================
with tab3:
    if "last_contract" not in st.session_state or "last_flags" not in st.session_state:
        st.info("💡 Please analyze a contract in the **Analyze Contract** tab first.")
    else:
        st.write("Select clauses you want removed or improved:")
        chosen = []
        for flag in st.session_state.last_flags:
            flag_key = f"flag_{hashlib.md5(flag.encode()).hexdigest()}"
            if st.checkbox(flag, key=flag_key):
                chosen.append(flag)

        if st.button("✍️ Generate Improved Version", key="btn_improve"):
            if not chosen:
                st.warning("Please select at least one clause to improve.")
            else:
                try:
                    with st.spinner("Generating improved contract..."):
                        improved = rewrite_contract(
                            st.session_state.last_contract,
                            chosen,
                            model_choice
                        )
                    st.session_state.improved_contract = improved
                except Exception as e:
                    st.error(f"Failed to generate improved contract: {e}")

        if "improved_contract" in st.session_state:
            st.success("✅ Contract improved successfully!")
            st.subheader("Improved Contract")
            st.text_area("", st.session_state.improved_contract, height=400)
            doc_file = create_doc(st.session_state.improved_contract)
            st.download_button(
                "⬇️ Download as Word Document",
                doc_file,
                file_name="Improved_Contract.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="btn_download"
            )

# =========================================================
# FOOTER
# =========================================================
st.markdown("""
<div class="dracu-footer">
    ⚖️ <strong>Dracu-Law</strong> &nbsp;·&nbsp; Powered by <strong>Groq LLMs</strong> &nbsp;·&nbsp; Built with <strong>Streamlit</strong><br>
    <span style="font-size:11px;color:#444;">AI-generated analysis is not a substitute for professional legal advice.</span>
</div>
""", unsafe_allow_html=True)